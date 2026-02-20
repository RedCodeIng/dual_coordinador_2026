import sys
import os
from docx import Document

# Add src to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../src')))

from utils.docs_generator import generate_document

def create_dummy_template(path):
    print(f"Creando plantilla dummy en: {path}")
    doc = Document()
    doc.add_heading('Carta de Aceptación', 0)
    p = doc.add_paragraph('Estimado/a ')
    p.add_run('{{ nombre_alumno }}').bold = True
    p.add_run(',')
    
    doc.add_paragraph('Por medio de la presente se confirma su aceptación en el programa Dual para el periodo {{ periodo }}.')
    
    doc.add_paragraph('Atentamente,')
    doc.add_paragraph('Coordinación Dual')
    
    doc.save(path)
    print(f"Plantilla dummy creada exitosamente.")

def test_generate_document():
    # El módulo docs_generator busca en src/templates/docs por defecto
    src_template_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../src/templates/docs'))
    if not os.path.exists(src_template_dir):
        print(f"Creando directorio de plantillas: {src_template_dir}")
        os.makedirs(src_template_dir)

    template_name = "test_template.docx"
    template_path = os.path.join(src_template_dir, template_name)
    
    output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'output'))
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_path = os.path.join(output_dir, "generated_doc.docx")
    
    # Crear plantilla si no existe
    if not os.path.exists(template_path):
        create_dummy_template(template_path)
    
    context = {
        "nombre_alumno": "Juan Pérez",
        "periodo": "Enero - Junio 2024"
    }
    
    print(f"Generando documento desde {template_name}...")
    success, message = generate_document(template_name, output_path, context)
    
    if success:
        print(f"✅ Documento generado exitosamente.")
        print(f"Verificar archivo en: {output_path}")
    else:
        print(f"❌ Error al generar documento: {message}")

if __name__ == "__main__":
    test_generate_document()
