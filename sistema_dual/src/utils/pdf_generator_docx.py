import os
import copy
import docx
from datetime import datetime
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate

def clone_and_replace(row, text):
    """
    Creates a deep copy of an existing row to ensure we preserve all
    MS Word cell layouts, grid spans, and borders.
    """
    new_tr = copy.deepcopy(row._tr)
    new_row = docx.table._Row(new_tr, row._parent)
    
    # clear all cells
    for cell in new_row.cells:
        for p in cell.paragraphs:
            p.text = ''
            
    # set text in first cell
    if len(new_row.cells) > 0:
        new_row.cells[0].paragraphs[0].text = text
        
    return new_tr

def generate_pdf_from_docx(template_name, context, output_path, template_path=None):
    """
    Generates a PDF from a DOCX template using docxtpl and docx2pdf.
    Dynamically injects {%tr for %} rows to multiply list elements,
    and inserts a page break before the signatures block.
    """
    try:
        if template_path and os.path.exists(template_path):
            t_path = template_path
        else:
            base_dir = os.path.join(os.path.dirname(__file__), '../templates/docs')
            t_path = os.path.join(base_dir, template_name)
            
        if not os.path.exists(t_path):
            return False, f"Plantilla DOCX no encontrada: {t_path}"

        # 2. Preprocess Template to inject dynamic rows and page break
        raw_doc = docx.Document(t_path)
        elaboraron_processed = False
        
        for t_idx, table in enumerate(raw_doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = ''.join(cell.text for cell in row.cells)
                
                if '{{ competencia }}' in row_text:
                    row._tr.addprevious(clone_and_replace(row, '{%tr for c in competencias_list %}'))
                    row._tr.addnext(clone_and_replace(row, '{%tr endfor %}'))
                    
                    processed_cells = set()
                    for cell in row.cells:
                        if cell in processed_cells: continue
                        processed_cells.add(cell)
                        for p in cell.paragraphs:
                            if '{{' in p.text:
                                p.text = p.text.replace('{{ loop_index }}', '{{ loop.index }}') \
                                               .replace('{{ competencia }}', '{{ c.competencia }}') \
                                               .replace('{{ asignatura }}', '{{ c.asignatura }}')
                                
                elif '{{ actividad }}' in row_text:
                    row._tr.addprevious(clone_and_replace(row, '{%tr for a in actividades_list %}'))
                    row._tr.addnext(clone_and_replace(row, '{%tr endfor %}'))
                    
                    processed_cells = set()
                    for cell in row.cells:
                        if cell in processed_cells: continue
                        processed_cells.add(cell)
                        for p in cell.paragraphs:
                            if '{{' in p.text:
                                p.text = p.text.replace('{{ loop_index }}', '{{ loop.index }}') \
                                               .replace('{{ actividad }}', '{{ a.actividad }}') \
                                               .replace('{{ horas }}', '{{ a.horas }}') \
                                               .replace('{{ evidencia }}', '{{ a.evidencia }}') \
                                               .replace('{{ lugar }}', '{{ a.lugar }}') \
                                               .replace('{{ ponderacion }}', '{{ a.ponderacion }}')

                # Check for ELABORARON (Signatures block) - Insert page break before the signature table
                if 'ELABORARON' in row_text and not elaboraron_processed:
                    processed_cells = set()
                    for cell in row.cells:
                        if cell in processed_cells: continue
                        processed_cells.add(cell)
                        for p in cell.paragraphs:
                            if 'ELABORARON' in p.text:
                                p.insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)
                                elaboraron_processed = True
                                break
                        if elaboraron_processed: break

        injected_path = output_path.replace('.pdf', '_INJECTED.docx')
        if not injected_path.endswith('.docx'): injected_path += '_INJECTED.docx'
        raw_doc.save(injected_path)

        # 3. Render Template using docxtpl on the injected file
        doc = DocxTemplate(injected_path)
        
        # docxtpl uses jinja2 under the hood, we just pass the context
        doc.render(context)
            
        # 4. Save filled DOCX temporarily
        temp_docx_path = output_path.replace('.pdf', '.docx')
        if not temp_docx_path.endswith('.docx'):
             temp_docx_path += '.docx'
             
        doc.save(temp_docx_path)
        
        # 6. Convert to PDF using docx2pdf (Requires MS Word installed)
        try:
            from docx2pdf import convert
            convert(temp_docx_path, output_path)
            
            # Clean up temporary DOCX files to save resources
            try:
                os.remove(injected_path)
                os.remove(temp_docx_path)
            except Exception as e:
                print(f"Cleanup warning: {e}")
            
            return True, f"Documento generado exitosamente en: {output_path}"
            
        except ImportError:
            # If docx2pdf fails or isn't available, we return the filled DOCX
            temp_pdf_path = output_path
            return True, f"DOCX rellenado exitosamente en: {temp_docx_path} (No se pudo convertir a PDF automáticamente)"
        except Exception as e:
            return True, f"DOCX guardado en {temp_docx_path}. Error al convertir a PDF: {e}"

    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, f"Error interactuando con DOCX: {str(e)}"
